#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def md_to_docx(md_file, docx_file):
    """将Markdown文件转换为Word文档"""
    # 创建Word文档
    doc = Document()
    
    # 设置样式
    styles = doc.styles
    
    # 标题1样式
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    
    # 标题2样式
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x29, 0x80, 0xb9)
    
    # 正文样式
    normal_style = styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.font.name = 'Microsoft YaHei'
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.readlines()
    
    # 处理每一行
    in_code_block = False
    code_content = []
    
    for line in md_content:
        line = line.rstrip()
        
        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                in_code_block = False
                
                # 添加代码块内容
                p = doc.add_paragraph()
                p.style = 'Normal'
                code_text = '\n'.join(code_content)
                p.add_run(code_text).font.name = 'Courier New'
                p.add_run().add_break()
                code_content = []
            else:
                # 开始代码块
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # 处理无序列表
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:])
            p.style = 'List Bullet'
        
        # 处理有序列表
        elif line.startswith('1. ') or (line and line[0].isdigit() and line[1:].startswith('. ')):
            p = doc.add_paragraph(line[line.find('.')+2:])
            p.style = 'List Number'
        
        # 处理普通段落
        elif line:
            if line.startswith('**') and line.endswith('**'):
                # 粗体标记
                p = doc.add_paragraph()
                p.add_run(line.strip('**')).bold = True
            else:
                p = doc.add_paragraph(line)
                p.style = 'Normal'
        
        # 空行
        else:
            doc.add_paragraph()
    
    # 保存文档
    doc.save(docx_file)
    print(f"已将 {md_file} 转换为 {docx_file}")

if __name__ == "__main__":
    md_file = "20250506网址导航软件添加锁定功能及修正开发记录.md"
    docx_file = "20250506网址导航软件添加锁定功能及修正开发记录.docx"
    
    md_to_docx(md_file, docx_file)